# !/usr/bin/python
# -*- coding: utf-8 -*-
# @time    : 2020/5/27 21:21
# @author  : Mo
# @function: text tools of common, 通用文本数据分析基础函数


from text_analysis.conf.path_log import logger
import json
import re
import os


__all__ = ["get_all_dirs_files",
           "get_dir_files",
           "delete_file",
           "save_json",
           "load_json"
           "docx_read",
           "txt_write",
           "txt_read",
           ]


def get_all_dirs_files(path_dir):
    """
        递归获取某个目录下的所有文件(所有层, 包括子目录)
    :param path_dir: str, like '/home/data'
    :return: list, like ['2020_01_08.txt']
    """
    path_files = []
    def get_path_files(path_dir):
        """
            递归函数, 获取某个目录下的所有文件
        :param path_dir: str, like '/home/data'
        :return: list, like ['2020_01_08.txt']
        """
        for root, dirs, files in os.walk(path_dir):
            for fi in files: # 递归的终止条件
                path_file = os.path.join(root, fi)
                path_files.append(path_file)
            for di in dirs:  # 语间目录便继续递归
                path_dir = os.path.join(root, di)
                get_path_files(path_dir)
    get_path_files(path_dir)
    return path_files


def get_dir_files(path_dir):
    """
        递归获取某个目录下的所有文件(单层)
    :param path_dir: str, like '/home/data'
    :return: list, like ['2019_12_5.txt']
    """

    def get_dir_files_func(file_list, dir_list, root_path=path_dir):
        """
            递归获取某个目录下的所有文件
        :param root_path: str, like '/home/data'
        :param file_list: list, like []
        :param dir_list: list, like []
        :return: None
        """
        # 获取该目录下所有的文件名称和目录名称
        dir_or_files = os.listdir(root_path)
        for dir_file in dir_or_files:
            # 获取目录或者文件的路径
            dir_file_path = os.path.join(root_path, dir_file)
            # 判断该路径为文件还是路径
            if os.path.isdir(dir_file_path):
                dir_list.append(dir_file_path)
                # 递归获取所有文件和目录的路径
                get_dir_files_func(dir_file_path, file_list, dir_list)
            else:
                file_list.append(dir_file_path)

    # 用来存放所有的文件路径
    _files = []
    # 用来存放所有的目录路径
    dir_list = []
    get_dir_files_func(_files, dir_list, path_dir)
    return _files


def delete_file(path):
    """
        删除一个目录下的所有文件
    :param path: str, dir path
    :return: None
    """
    for i in os.listdir(path):
        # 取文件或者目录的绝对路径
        path_children = os.path.join(path, i)
        if os.path.isfile(path_children):
            if path_children.endswith(".h5") or path_children.endswith(".json") \
                    or "events" in path_children or "trace" in path_children:
                os.remove(path_children)
        else:# 递归, 删除目录下的所有文件
            delete_file(path_children)


def txt_read(path_file, encode_type="utf-8"):
    """
        读取txt文件，默认utf8格式, 不能有空行
    :param file_path: str, 文件路径
    :param encode_type: str, 编码格式
    :return: list
    """
    lines = []
    try:
        file = open(path_file, "r", encoding=encode_type)
        lines = file.readlines()
        file.close()
    except Exception as e:
        logger.info(str(e))
    finally:
        return lines


def txt_write(list_line, file_path, type='w', encode_type='utf-8'):
    """
      txt写入list文件
    :param listLine:list, list文件，写入要带"\n" 
    :param filePath:str, 写入文件的路径
    :param type: str, 写入类型, w, a等
    :param encode_type: 
    :return: 
    """
    try:
        file = open(file_path, type, encoding=encode_type)
        file.writelines(list_line)
        file.close()
    except Exception as e:
        logger.info(str(e))


def save_json(json_lines, json_path, encoding='utf-8', indent=4):
    """
      保存json，
    :param json_lines: json 
    :param path: str
    :return: None
    """
    with open(json_path, 'w', encoding=encoding) as fj:
        fj.write(json.dumps(json_lines, ensure_ascii=False, indent=indent))
    fj.close()


def load_json(path, encoding="utf-8"):
    """
      获取json, json存储为[{}]格式, like [{'大漠帝国':132}]
    :param path: str
    :return: json
    """
    with open(path, 'r', encoding=encoding) as fj:
        model_json = json.load(fj)
    return model_json


def docx_read(path):
    """读取word文档的段落数据text/table, docx
    read corpus from docx
    Args:
        path: String, path/origin text,  eg. "叉尾斗鱼介绍.docx"
    Returns:
        passages: List<tuple>, 文章段落
    """

    def iter_tables(block_item_container):
        """Recursively generate all tables in `block_item_container`."""
        for t in block_item_container.tables:
            yield t
            for row in t.rows:
                for cell in row.cells:
                    yield from iter_tables(cell)
    passages = []
    try:
        import docx
        docx_temp = docx.Document(path)
        # 文本
        for p in docx_temp.paragraphs:
            if p.text.strip():
                passages.append(p.text.strip()+"\n")
        # 表格
        for t in iter_tables(docx_temp):
            table = t
            df = [["" for i in range(len(table.columns))] for j in range(len(table.rows))]
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if cell.text:
                        df[i][j] = cell.text.replace("\n", "")
            df = [" ".join(dfi).strip() + "\n" for dfi in df]
            passages += df
    except Exception as e:
        logger.info(str(e))
    return passages


def pdf_read(path):
    """读取pdf文档的段落数据text/table, docx
    read corpus from docx
    Args:
        path: String, path/origin text,  eg. "叉尾斗鱼介绍.pdf"
    Returns:
        passages: List<tuple>, 文章段落
    """
    passages = []
    try:
        from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
        from pdfminer.converter import TextConverter
        from pdfminer.pdfdocument import PDFDocument
        from pdfminer.pdfparser import PDFParser
        from pdfminer.pdfpage import PDFPage
        from pdfminer.layout import LAParams
        from io import StringIO

        output_string = StringIO()
        with open(path, "rb") as prb:
            parser = PDFParser(prb)
            doc = PDFDocument(parser)
            rsrcmgr = PDFResourceManager()
            device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
            interpreter = PDFPageInterpreter(rsrcmgr, device)
            for page in PDFPage.create_pages(doc):
                interpreter.process_page(page)

        passages = output_string.getvalue()
    except Exception as e:
        logger.info(str(e))
    return passages


def docx_extract_color(path, highligt_colors=["YELLOW"]):
    """抽取docx文档中高亮的text, 没有table
    docx extract color
    Args:
        path: String, path/origin text,  eg. "叉尾斗鱼介绍.docx"
        highligt_colors: List, enum/highligt colors,  eg. ["YELLOW", "RED"]
    Returns:
        passages: List<tuple>, 文章段落
    """
    passages = []
    try:
        from docx.enum.text import WD_COLOR_INDEX
        import docx
        wd_colors_dict = {wcim.name:wcim.value for wcim in WD_COLOR_INDEX.__members__[2:]}
        wd_colors_select = [wd_colors_dict.get(hc.upper(), "") for hc in highligt_colors
                                 if wd_colors_dict.get(hc.upper(), "")]
        """
        # AUTO = 'default'
        BLACK = 'black'
        BLUE = 'blue'
        BRIGHTGREEN = 'green'
        DARKBLUE = 'darkBlue'
        DARKRED = 'darkRed'
        DARKYELLOW = 'darkYellow'
        GRAY25 = 'lightGray'
        GRAY50 = 'darkGray'
        GREEN = 'darkGreen'
        PINK = 'magenta'
        RED = 'red'
        TEAL = 'darkCyan'
        TURQUOISE = 'cyan'
        VOILET = 'darkMagenta'
        WHITE = 'white'
        YELLOW = 'yellow'
        """
        document = docx.Document(path)
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                if run.font.highlight_color and run.font.highlight_color in wd_colors_select:
                    passages.append(run.text.strip()+"\n")
                elif not highligt_colors and not run.font.highlight_color:
                    passages.append(run.text.strip()+"\n")
    except Exception as e:
        logger.info(str(e))
    return passages


if __name__ == '__main__':
    path = "../data/corpus/原始语料/CCKS2021技术评测-面向中文电子病历的医疗实体及事件抽取.docx"
    passages = docx_read(path)
    for p in passages:
        print(p)

