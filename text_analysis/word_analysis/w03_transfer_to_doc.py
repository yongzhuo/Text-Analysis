# !/usr/bin/python
# -*- coding: utf-8 -*-
# @time    : 2021/10/23 10:14
# @author  : Mo
# @function: cut word txt, 将txt文档切词


from text_analysis.utils.text_common import docx_read, txt_read, txt_write, get_all_dirs_files
from text_analysis.conf.path_log import logger
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
from docx import Document
from tqdm import tqdm
import jieba
import os
import re


def write_txt_to_docx(path_out, texts, keywords_field=None, keywords_common=None, keywords_newword=None):
    """将txt写入docx文档
        write txt to docx
        Args:
            path_out: String, path_out/origin text,  eg. "切词.docx"
            texts: List, text,  eg. ["你/是/谁/呀"]
            keywords_field: List, text,  eg. ["金融社保卡"]
            keywords_common: List, text,  eg. ["北京"]
        Returns:
            None
    """
    def extract_string(text, regx="###"):
        """抽取string中间的特定字符串
        extract string
        Args:
            text: String, path_in_dicr/origin text,  eg. "dir_txt"
            regx: String, path_save_dir/save text,  eg. "dir_txt_cut"
        Returns:
            text_ext: List<str>
        """
        pattern = r"{}(.*?){}".format(regx, regx)
        text_ext = re.findall(pattern, text)
        return text_ext

    document = Document()
    count = 0
    for passage in tqdm(texts):
        if not passage.strip():
            continue
        count += 1
        if count % 2 == 0:
            document.add_paragraph("\n")
            document.add_paragraph(passage.replace("/", ""))

        document.add_paragraph("【")
        words = passage.split("/")

        """
        AUTO = 'default'
        BLACK = 'black'
        BLUE = 'blue'
        BRIGHTGREEN = 'green'
        DARKBLUE = 'darkBlue'
        DARKRED = 'darkRed'
        DARKYELLOW = 'darkYellow'
        GRAY25 = 'lightGray'
        GRAY50 = 'darkGray'
        GREEN = 'darkGreen'
        PINK = 'magenta'
        RED = 'red'
        TEAL = 'darkCyan'
        TURQUOISE = 'cyan'
        VOILET = 'darkMagenta'
        WHITE = 'white'
        YELLOW = 'yellow'
        """

        par = document.paragraphs[-1]
        for w in words:
            par.add_run(text="/")
            par.add_run(text=w)
            run = par.runs[-1]
            # run.font.style = "宋体"  # 文字大小
            run.font.size = Pt(10)  # 文字大小
            if len(w) == 1:
                run.font.highlight_color = WD_COLOR_INDEX.DARK_YELLOW
            elif w in keywords_newword:
                run.font.highlight_color = WD_COLOR_INDEX.BRIGHTGREEN
            elif w in keywords_field:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif w in keywords_common:
                run.font.highlight_color = WD_COLOR_INDEX.GREEN
            else:
                run.font.highlight_color = None  # WD_COLOR_INDEX.RED
        document.save(path_out)



def transfer_txt_to_docx(path_in_dir, path_save_dir=None, path_keywords_field="keywords/keywords_field.txt",
                    path_keywords_common="keywords/keywords_common.txt",
                    path_keywords_newword="keywords/keywords_newword.txt"):
    """将txt写入docx文档, 并标颜色
        write txt to docx
        Args:
            path_in_dir: String, path_in_dicr/origin text,  eg. "dir_txt"
            path_save_dir: String, path_save_dir/save text,  eg. "dir_txt_cut"
            path_keywords_field: String, path_keywords_field/keywords file,  eg. "keywords_field.txt"
            path_keywords_common: String, path_keywords_common/keywords file,  eg. "keywords_common.txt"
            path_keywords_newword: String, path_keywords_newword/keywords file,  eg. "keywords_newword.txt"
        Returns:
            None
    """
    # 关键词
    keywords_newword = open(path_keywords_newword, "r", encoding="utf-8").readlines()
    keywords_common = open(path_keywords_common, "r", encoding="utf-8").readlines()
    keywords_field = open(path_keywords_field, "r", encoding="utf-8").readlines()
    keywords_field = [kf.strip() for kf in keywords_field]
    for kf in keywords_field:
        jieba.add_word(kf, 12000)
    # 创建存储的目录名
    if path_save_dir:
        if not os.path.exists(path_save_dir):
            os.mkdir(path_save_dir)
    else:  # 如果没有就存在原目录
        path_save_dir = path_in_dir

    files = get_all_dirs_files(path_in_dir)
    for file in files:
        if ".切词.txt" in file:
            print(file)
            texts = open(file, "r", encoding="utf-8").readlines()
            # texts = ["/".join(jieba.lcut(t, HMM=True)) for t in texts if t.strip()]
            # 最后一层文件名
            file_name = file.replace(path_in_dir, "").strip("/").strip("\\")
            file_name = file_name.replace(file_name.split(".")[-1], "docx")
            path_save_file = os.path.join(path_save_dir, file_name)
            write_txt_to_docx(path_save_file, texts, keywords_field, keywords_common, keywords_newword)
    ta = 0


def os_main(path_in_dir, path_save_dir):
    """ 转为doc文件 """
    path_keywords_common = "keywords/keywords_common.txt"
    path_keywords_field = "keywords/keywords_field.txt"

    path_in_dir = os.path.join(path_save_dir, "中文分词")
    path_save_dir = os.path.join(path_save_dir, "中文分词docx")

    # path_in_dir = "../data/corpus/分析结果/中文分词"
    # path_save_dir = "../data/corpus/分析结果/中文分词docx"
    transfer_txt_to_docx(path_in_dir, path_save_dir, path_keywords_field, path_keywords_common)
    logger.info("w03_transfer_to_doc.py ok")
    ta = 0


if __name__ == '__main__':

    path_keywords_common = "keywords/keywords_common.txt"
    path_keywords_field = "keywords/keywords_field.txt"
    path_in_dir = "../data/corpus/分析结果/中文分词"
    path_save_dir = "../data/corpus/分析结果/中文分词docx"
    transfer_txt_to_docx(path_in_dir, path_save_dir, path_keywords_field, path_keywords_common)
    ee = 0

